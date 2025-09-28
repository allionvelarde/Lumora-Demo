# app.py - Lumora (full app with polished blue landing UI + functional pages)
import streamlit as st
import sqlite3
import os
import hashlib
import re
import random
import string
import json
from datetime import datetime
from io import BytesIO

# optional AI / file libs
import openai
from PyPDF2 import PdfReader
from docx import Document

# ---------- CONFIG ----------
DB_PATH = "lumora.db"
NAME_PATTERN = r"^[A-Z][a-zA-Z]+ [A-Z]\. [A-Z][a-zA-Z]+$"  # "First M. Last" e.g., Juan D. Cruz

# Set OpenAI API key from env if available
openai.api_key = os.getenv("OPENAI_API_KEY", None)

# ---------- DB helpers & migration ----------
def get_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def init_db():
    conn = get_conn()
    c = conn.cursor()

    # users: ensure password column exists (migrate if needed)
    c.execute("""CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY,
        name TEXT UNIQUE,
        role TEXT,
        password TEXT
    )""")
    # survey_results
    c.execute("""CREATE TABLE IF NOT EXISTS survey_results (
        id INTEGER PRIMARY KEY,
        user_id INTEGER,
        mi_scores TEXT,
        dominant_mi TEXT,
        created_at TEXT
    )""")
    # lessons (personal/generated)
    c.execute("""CREATE TABLE IF NOT EXISTS lessons (
        id INTEGER PRIMARY KEY,
        user_id INTEGER,
        topic TEXT,
        subject TEXT,
        mi_type TEXT,
        content TEXT,
        created_at TEXT
    )""")
    # classes
    c.execute("""CREATE TABLE IF NOT EXISTS classes (
        id INTEGER PRIMARY KEY,
        teacher_id INTEGER,
        name TEXT,
        code TEXT UNIQUE,
        created_at TEXT
    )""")
    # class_members
    c.execute("""CREATE TABLE IF NOT EXISTS class_members (
        id INTEGER PRIMARY KEY,
        class_id INTEGER,
        student_id INTEGER,
        joined_at TEXT
    )""")
    # class_lessons (tailored copies for each student)
    c.execute("""CREATE TABLE IF NOT EXISTS class_lessons (
        id INTEGER PRIMARY KEY,
        class_id INTEGER,
        student_id INTEGER,
        topic TEXT,
        subject TEXT,
        mi_type TEXT,
        content TEXT,
        created_at TEXT
    )""")

    conn.commit()
    conn.close()

def hash_pw(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def user_exists(name: str) -> bool:
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id FROM users WHERE name=?", (name,))
    r = c.fetchone(); conn.close()
    return bool(r)

def create_user(name: str, role: str, pw: str):
    conn = get_conn(); c = conn.cursor()
    c.execute("INSERT INTO users (name, role, password) VALUES (?, ?, ?)", (name, role, hash_pw(pw)))
    conn.commit(); uid = c.lastrowid; conn.close()
    return uid

def validate_user(name: str, pw: str):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id, role, password FROM users WHERE name=?", (name,))
    row = c.fetchone(); conn.close()
    if row and row[2] == hash_pw(pw):
        return row[0], row[1]
    return None, None

def save_survey(user_id, mi_scores: dict, dominant: str):
    conn = get_conn(); c = conn.cursor()
    now = datetime.now().isoformat()
    c.execute("INSERT INTO survey_results (user_id, mi_scores, dominant_mi, created_at) VALUES (?,?,?,?)",
              (user_id, json.dumps(mi_scores), dominant, now))
    conn.commit(); conn.close()

def get_latest_dominant_mi(user_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT dominant_mi FROM survey_results WHERE user_id=? ORDER BY id DESC LIMIT 1", (user_id,))
    row = c.fetchone(); conn.close()
    return row[0] if row else None

def save_lesson(user_id, topic, subject, mi_type, content):
    conn = get_conn(); c = conn.cursor()
    now = datetime.now().isoformat()
    c.execute("INSERT INTO lessons (user_id, topic, subject, mi_type, content, created_at) VALUES (?,?,?,?,?,?)",
              (user_id, topic, subject, mi_type, content, now))
    lid = c.lastrowid; conn.commit(); conn.close()
    return lid

def get_lessons_by_user(user_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id, topic, subject, mi_type, content, created_at FROM lessons WHERE user_id=? ORDER BY id DESC", (user_id,))
    rows = c.fetchall(); conn.close()
    return rows

def get_all_lessons():
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id, topic, subject, mi_type, content, created_at FROM lessons ORDER BY id DESC")
    rows = c.fetchall(); conn.close()
    return rows

def create_class(teacher_id, name):
    code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
    conn = get_conn(); c = conn.cursor()
    now = datetime.now().isoformat()
    c.execute("INSERT INTO classes (teacher_id, name, code, created_at) VALUES (?,?,?,?)",
              (teacher_id, name, code, now))
    conn.commit(); conn.close()
    return code

def get_teacher_classes(teacher_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id, name, code FROM classes WHERE teacher_id=?", (teacher_id,))
    rows = c.fetchall(); conn.close()
    return rows

def join_class_by_code(student_id, code):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id FROM classes WHERE code=?", (code,))
    row = c.fetchone()
    if not row:
        conn.close()
        return False, "Invalid class code."
    class_id = row[0]
    # check already member
    c.execute("SELECT id FROM class_members WHERE class_id=? AND student_id=?", (class_id, student_id))
    if c.fetchone():
        conn.close()
        return True, "Already a member."
    now = datetime.now().isoformat()
    c.execute("INSERT INTO class_members (class_id, student_id, joined_at) VALUES (?,?,?)", (class_id, student_id, now))
    conn.commit(); conn.close()
    return True, "Joined successfully."

def get_student_classes(student_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("""
        SELECT c.id, c.name, c.code
        FROM classes c
        JOIN class_members m ON c.id = m.class_id
        WHERE m.student_id = ?
    """, (student_id,))
    rows = c.fetchall(); conn.close()
    return rows

def get_students_in_class(class_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("""
        SELECT u.id, u.name
        FROM users u
        JOIN class_members m ON u.id = m.student_id
        WHERE m.class_id = ?
    """, (class_id,))
    rows = c.fetchall(); conn.close()
    return rows

def save_class_lesson(class_id, student_id, topic, subject, mi_type, content):
    conn = get_conn(); c = conn.cursor()
    now = datetime.now().isoformat()
    c.execute("""INSERT INTO class_lessons (class_id, student_id, topic, subject, mi_type, content, created_at)
                 VALUES (?,?,?,?,?,?,?)""", (class_id, student_id, topic, subject, mi_type, content, now))
    conn.commit(); conn.close()

def get_class_lessons_for_student(student_id, class_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("""SELECT topic, subject, mi_type, content, created_at
                 FROM class_lessons
                 WHERE student_id=? AND class_id=?
                 ORDER BY id DESC""", (student_id, class_id))
    rows = c.fetchall(); conn.close()
    return rows

# initialize DB
init_db()

# ---------- AI helper (guarded) ----------
def generate_tailored_lesson(text: str, mi_type: str) -> str:
    """
    Calls OpenAI to rewrite/extract/structure the lesson tailored to mi_type.
    If OpenAI is not configured or call fails, returns fallback text and raises no raw errors to student.
    """
    prompt = f"Rewrite and structure this lesson content for a student with {mi_type} intelligence. " \
             f"Keep it concise, give 3-4 key points and 2 practice questions.\n\nOriginal content:\n{text}"
    if not openai.api_key:
        # No API key configured -> return fallback structured text (simple transformation)
        # Provide a small local "tailoring" fallback so students still receive something useful.
        fallback = f"[Tailored for {mi_type}] (No OpenAI key configured)\n\n"
        fallback += "Key points:\n- " + (text[:200].replace("\n", " ") if text else "Main idea placeholder") + "\n\nPractice:\n1) ...\n2) ..."
        return fallback

    try:
        resp = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=700,
            temperature=0.7
        )
        out = resp["choices"][0]["message"]["content"].strip()
        return out
    except Exception as e:
        # log error in Streamlit UI (visible to teacher) but return a safe fallback
        st.error(f"AI generation failed: {e}")
        fallback = f"[Tailored for {mi_type}] (AI fallback)\n\n" + (text[:400] if text else "Content not available.")
        return fallback

# ---------- UI / session init ----------
st.set_page_config(page_title="Lumora", layout="wide", initial_sidebar_state="collapsed")
if "page" not in st.session_state: st.session_state.page = "auth"  # 'auth' or 'app'
for k in ("user_id", "user_name", "user_role", "logged_in"):
    if k not in st.session_state:
        st.session_state[k] = None

# ---------- TOP polished landing (login/signup) UI ----------
# We'll show this when not logged in; inside app pages remain simple-blue
def landing_ui():
    # animated background + form/card CSS
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;700&display=swap');

    /* Base app background */
    .stApp {
        background: linear-gradient(135deg, #07103a 0%, #0f2a66 40%, #2563eb 100%);
        overflow: hidden;
        font-family: 'Poppins', sans-serif;
        color: #e6eef8;
        min-height: 100vh;
    }

    /* floating blurred glow lights behind everything (fixed so it doesn't affect layout) */
    body::before, body::after {
        content: "";
        position: fixed;
        width: 600px;
        height: 600px;
        border-radius: 50%;
        filter: blur(160px);
        opacity: 0.35;
        z-index: -9999;
        animation: float 24s infinite ease-in-out alternate;
    }
    body::before {
        top: -200px; left: -200px;
        background: radial-gradient(circle at center, rgba(59,130,246,0.8), rgba(14,165,233,0.4), transparent 60%);
        animation-delay: 0s;
    }
    body::after {
        bottom: -200px; right: -200px;
        background: radial-gradient(circle at center, rgba(99,102,241,0.8), rgba(236,72,153,0.35), transparent 60%);
        animation-delay: 8s;
    }
    @keyframes float {
        0%   { transform: translate(0, 0) scale(1); }
        50%  { transform: translate(120px, -80px) scale(1.05); }
        100% { transform: translate(-80px, 100px) scale(0.95); }
    }

    /* static hero glass card (pure HTML block) */
    .hero-card {
        background: rgba(255,255,255,0.06);
        border-radius: 16px;
        padding: 36px;
        margin: 30px auto 18px auto;
        max-width: 1100px;
        box-shadow: 0 8px 40px rgba(3,37,65,0.5);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.06);
    }

    .title-gradient {
        background: linear-gradient(90deg, #9ec7ff, #60a5fa, #3b82f6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
        font-size: 44px;
        margin-bottom: 6px;
    }
    .muted { color: #cfe7ff; margin-bottom: 12px; font-size: 16px; }

    /* Style Streamlit forms to look like glass cards (this is the key part) */
    form.stBlock, form.stForm {
        background: rgba(255,255,255,0.06) !important;
        border-radius: 14px !important;
        padding: 22px !important;
        margin: 14px auto !important;
        max-width: 420px;
        box-shadow: 0 8px 30px rgba(3,37,65,0.45) !important;
        border: 1px solid rgba(255,255,255,0.06) !important;
        backdrop-filter: blur(8px);
    }

    /* style inputs and buttons inside Streamlit form */
    .stTextInput>div>div>input, .stPasswordInput>div>div>input, textarea {
        background: rgba(255,255,255,0.03) !important;
        color: #e6eef8 !important;
        border-radius: 8px !important;
        border: 1px solid rgba(255,255,255,0.06) !important;
        padding: 10px !important;
    }
    div.stButton > button {
        background: linear-gradient(90deg,#3b82f6,#06b6d4) !important;
        color:white !important;
        border-radius:10px !important;
        padding:8px 16px !important;
        font-weight:700 !important;
        box-shadow: 0 8px 24px rgba(59,130,246,0.28) !important;
    }
    div.stButton > button:hover { transform: scale(1.03); }

    /* center helper */
    .center-wrap { display:flex; justify-content:center; align-items:center; flex-direction:column; }
    </style>
    """, unsafe_allow_html=True)

    # HERO (pure HTML block)
    st.markdown(
        """
        <div class="hero-card">
            <div class="title-gradient">Lumora</div>
            <div class="muted">Personalized lessons for every intelligence</div>
            <div style="color:#dbeafe; font-size:16px; line-height:1.5;">
                Generate lessons, join classes, and receive AI-tailored materials based on your dominant multiple intelligence.
                <br><br>
                For demo: sign up as Student or Teacher, create/join classes, upload lessons — AI will tailor content per student.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # LOGIN / SIGNUP form (uses st.form so the form container is stylable)
    st.markdown("<div class='center-wrap'>", unsafe_allow_html=True)
    mode = st.radio("Choose action:", ("Log in", "Sign up"), index=0, horizontal=True, key="auth_mode_ui")

    if mode == "Sign up":
        with st.form("signup_form"):
            st.subheader("Create an account")
            name = st.text_input("Full name (First M. Last)", key="signup_name_ui")
            role = st.selectbox("Role", ("student", "teacher"), key="signup_role_ui")
            pw = st.text_input("Password", type="password", key="signup_pw_ui")
            pw2 = st.text_input("Confirm password", type="password", key="signup_confirm_ui")
            submitted = st.form_submit_button("Create account")
            if submitted:
                if not name or not pw:
                    st.error("Please fill name and password.")
                elif not re.match(NAME_PATTERN, name.strip()):
                    st.error("Name must be like: Juan D. Cruz")
                elif pw != pw2:
                    st.error("Passwords do not match.")
                elif user_exists(name.strip()):
                    st.error("User already exists. Please log in.")
                else:
                    uid = create_user(name.strip(), role, pw)
                    if uid:
                        st.success("Account created. Please switch to 'Log in' to sign in.")

    else:
        with st.form("login_form"):
            st.subheader("Log in")
            name = st.text_input("Full name (First M. Last)", key="login_name_ui")
            pw = st.text_input("Password", type="password", key="login_pw_ui")
            submitted = st.form_submit_button("Log in")
            if submitted:
                if not name or not pw:
                    st.error("Enter name and password.")
                else:
                    uid, role = validate_user(name.strip(), pw)
                    if uid:
                        st.success("Login successful.")
                        st.session_state.user_id = uid
                        st.session_state.user_name = name.strip()
                        st.session_state.user_role = role
                        st.session_state.logged_in = True
                        st.session_state.page = "app"
                        st.experimental_rerun()
                    else:
                        st.error("Invalid credentials or name format.")
    st.markdown("</div>", unsafe_allow_html=True)

# ---------- Simple, clean inside pages styling helper ----------
def inside_style_css():
    st.markdown("""
    <style>
    .inside-title { color:#dbeafe; font-weight:800; }
    .accent { color:#bfdbfe; }
    .card-simple { background: rgba(255,255,255,0.03); padding:12px; border-radius:10px; border:1px solid rgba(255,255,255,0.05); }
    .small-muted { color:#c7d2fe; font-size:0.9rem; }
    </style>
    """, unsafe_allow_html=True)

# ---------- App pages (simple + blue accents) ----------
def student_pages():
    inside_style_css()
    st.sidebar.title(f"{st.session_state.user_name} (Student)")
    if st.sidebar.button("Log out"):
        for k in ("user_id","user_name","user_role","logged_in","page"):
            if k in st.session_state: del st.session_state[k]
        st.session_state.page = "auth"
        st.experimental_rerun()

    page = st.sidebar.radio("Navigate", ("MI Survey", "Lesson Generator", "Lesson Library", "My Classes"))

    if page == "MI Survey":
        st.header("Multiple Intelligences Survey", anchor=None)
        st.write("Rate statements from 1 (Disagree) to 5 (Agree).")
        questions = {
            "Linguistic": ["I enjoy reading and writing.", "I can explain ideas clearly in words.", "I like word games.", "I often write notes or stories.", "I remember better when I take notes."],
            "Logical-Mathematical": ["I like solving puzzles.", "I find patterns easily.", "I enjoy experiments.", "I can calculate quickly.", "I like data and logic."],
            "Visual-Spatial": ["I learn with pictures.", "I visualize objects easily.", "I enjoy drawing.", "I can read maps/directions.", "I imagine clearly."],
            "Bodily-Kinesthetic": ["I enjoy hands-on activities.", "I remember when I move or act.", "I am good at sports.", "I like building things.", "I learn by touching/doing."],
            "Musical": ["I remember things via songs.", "I notice sound patterns.", "I enjoy music.", "I like rhythms/melodies.", "I study better with music."],
            "Interpersonal": ["I enjoy group activities.", "I understand people's feelings.", "I like teamwork.", "I talk through problems with others.", "I learn from discussions."],
            "Intrapersonal": ["I reflect on my own learning.", "I set personal goals.", "I know my strengths/weaknesses.", "I learn independently.", "I enjoy journaling."]
        }
        responses = {}
        for mi, qs in questions.items():
            st.subheader(mi)
            scores = []
            cols = st.columns(2)
            for i, q in enumerate(qs):
                col = cols[i % 2]
                scores.append(col.slider(q, 1, 5, 3, key=f"{mi}_{i}"))
            responses[mi] = scores
        if st.button("Submit Survey"):
            mi_scores = {mi: sum(vals) for mi, vals in responses.items()}
            dominant = max(mi_scores, key=mi_scores.get)
            save_survey(st.session_state.user_id, mi_scores, dominant)
            st.success(f"Survey saved. Dominant intelligence: {dominant}")
            st.json(mi_scores)

    elif page == "Lesson Generator":
        st.header("Lesson Generator")
        st.markdown("<div class='card-simple'>", unsafe_allow_html=True)
        topic = st.text_input("Lesson topic", key="gen_topic")
        subject = st.selectbox("Subject", ["Mathematics", "Science", "English", "Other"], key="gen_subject")
        uploaded = st.file_uploader("Optional: upload PDF or DOCX", type=["pdf", "docx"], key="gen_upload")
        content = ""
        if uploaded:
            try:
                if uploaded.name.lower().endswith(".pdf"):
                    pdf = PdfReader(uploaded)
                    content = "\n".join([p.extract_text() or "" for p in pdf.pages])
                else:
                    doc = Document(uploaded); content = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                st.error(f"Could not read uploaded file: {e}")
        if st.button("Generate & Save"):
            if not topic and not content:
                st.error("Provide a topic or upload a file.")
            else:
                dominant = get_latest_dominant_mi(st.session_state.user_id) or "General"
                source_text = content if content else topic
                lesson_text = generate_tailored_lesson(source_text, dominant)
                save_lesson(st.session_state.user_id, topic or "Generated Lesson", subject, dominant, lesson_text)
                st.success("Lesson generated and saved to your library.")
                st.write(lesson_text)
        st.markdown("</div>", unsafe_allow_html=True)

    elif page == "Lesson Library":
        st.header("Your Lessons")
        lessons = get_lessons_by_user(st.session_state.user_id)
        if not lessons:
            st.info("No lessons yet. Generate one or ask your teacher to assign one.")
        else:
            for lid, topic, subject, mi_type, content, created in lessons:
                st.subheader(f"{topic} — {subject}")
                st.write(content)
                st.caption(f"Saved: {created}")

    elif page == "My Classes":
        st.header("My Classes")
        st.markdown("<div class='card-simple'>", unsafe_allow_html=True)
        code = st.text_input("Join class by code", key="join_code")
        if st.button("Join"):
            ok, msg = join_class_by_code(st.session_state.user_id, code.strip())
            if ok:
                st.success(msg)
            else:
                st.error(msg)
        st.markdown("</div>", unsafe_allow_html=True)

        classes = get_student_classes(st.session_state.user_id)
        if not classes:
            st.info("You haven't joined any classes yet.")
        else:
            for cid, cname, ccode in classes:
                st.subheader(f"{cname} (Code: {ccode})")
                lessons = get_class_lessons_for_student(st.session_state.user_id, cid)
                if not lessons:
                    st.info("No lessons assigned to this class yet.")
                else:
                    for topic, subject, mi_type, content, created in lessons:
                        st.markdown(f"**{topic}** — {subject}  \n*Tailored for: {mi_type}*\n")
                        st.write(content)
                        st.caption(f"Assigned: {created}")

# ---------- Teacher pages ----------
def teacher_pages():
    inside_style_css()
    st.sidebar.title(f"{st.session_state.user_name} (Teacher)")
    if st.sidebar.button("Log out"):
        for k in ("user_id","user_name","user_role","logged_in","page"):
            if k in st.session_state: del st.session_state[k]
        st.session_state.page = "auth"
        st.experimental_rerun()

    page = st.sidebar.radio("Navigate", ("Create Class", "Assign Lesson", "My Classes"))

    if page == "Create Class":
        st.header("Create Class")
        cname = st.text_input("Class name")
        if st.button("Create"):
            if not cname:
                st.error("Enter a class name.")
            else:
                code = create_class(st.session_state.user_id, cname)
                st.success(f"Class created. Share this code with students: {code}")

    elif page == "Assign Lesson":
        st.header("Assign Lesson (Upload PDF/DOCX)")
        classes = get_teacher_classes(st.session_state.user_id)
        if not classes:
            st.info("You have no classes. Create one first.")
        else:
            class_map = {c[0]: f"{c[1]} (Code {c[2]})" for c in classes}
            class_id = st.selectbox("Select class", list(class_map.keys()), format_func=lambda k: class_map[k])
            topic = st.text_input("Lesson topic", key="assign_topic")
            subject = st.text_input("Subject", key="assign_subject")
            uploaded = st.file_uploader("Upload PDF or DOCX to assign", type=["pdf", "docx"], key="assign_upload")
            # Inform if OpenAI key missing
            if not openai.api_key:
                st.warning("OpenAI API key not configured. Tailoring will use local fallback (limited).")
            if st.button("Assign to class"):
                if not uploaded:
                    st.error("Upload a file to assign.")
                else:
                    # extract text
                    try:
                        if uploaded.name.lower().endswith(".pdf"):
                            pdf = PdfReader(uploaded)
                            text = "\n".join([p.extract_text() or "" for p in pdf.pages])
                        else:
                            doc = Document(uploaded)
                            text = "\n".join([p.text for p in doc.paragraphs])
                    except Exception as e:
                        st.error(f"Could not read file: {e}")
                        text = ""
                    if not text.strip():
                        st.error("No extractable text found in the uploaded file.")
                    else:
                        students = get_students_in_class(class_id)
                        if not students:
                            st.info("No students in this class to assign to.")
                        else:
                            progress = st.progress(0)
                            total = len(students)
                            for i, (sid, sname) in enumerate(students, start=1):
                                dominant = get_latest_dominant_mi(sid) or "General"
                                tailored = generate_tailored_lesson(text, dominant)
                                save_class_lesson(class_id, sid, topic or "Uploaded Lesson", subject or "General", dominant, tailored)
                                progress.progress(int(i/total * 100))
                            st.success("Assigned tailored lessons to all students in class.")

    elif page == "My Classes":
        st.header("My Classes & Students")
        classes = get_teacher_classes(st.session_state.user_id)
        if not classes:
            st.info("No classes yet.")
        else:
            for cid, name, code in classes:
                st.subheader(f"{name} (Code: {code})")
                students = get_students_in_class(cid)
                if not students:
                    st.info("No students enrolled.")
                else:
                    st.write("Students:")
                    for sid, sname in students:
                        st.write(f"- {sname}")
                st.markdown("---")

# ---------- App runner ----------
def main():
    if not st.session_state.get("logged_in"):
        landing_ui()
        return

    # inside app
    if st.session_state.user_role == "student":
        student_pages()
    else:
        teacher_pages()

if __name__ == "__main__":
    main()
